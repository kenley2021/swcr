# -*- coding: utf-8 -*-
import logging
import importlib_resources
import codecs
from os.path import abspath
from os import scandir

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dataclasses import dataclass
import argparse

logger = logging.getLogger(__name__)

DEFAULT_INDIRS = ['.']
DEFAULT_EXTS = ['c', 'h']
DEFAULT_COMMENT_CHARS = ['/*', '*', '*/', '//']

def del_slash(dirs):
    return [dir_[:-1] if dir_[-1] == '/' else dir_ for dir_ in dirs]

class CodeFinder(object):
    def __init__(self, exts=None):
        self.exts = exts if exts else DEFAULT_EXTS

    @staticmethod
    def is_hidden_file(file):
        return file[0] == '.'

    @staticmethod
    def should_be_excluded(file, excludes = None):
        if not excludes:
            return False
        if not isinstance(excludes, list):
            excludes = [excludes]
        return any(file.startswith(exclude) for exclude in excludes)

    def is_code(self, file):
        is_code_file = any(file.endswith(ext) for ext in self.exts)
        return is_code_file

    def find(self, indir, excludes = None):
        files = []
        for entry in scandir(indir):
            entry_name = entry.name
            entry_path = abspath(entry.path)
            if self.is_hidden_file(entry_name) or self.should_be_excluded(entry_path, excludes):
                continue
            if entry.is_file():
                if self.is_code(entry_name):
                    files.append(entry_path)
            else:
                files.extend(self.find(entry_path, excludes=excludes))
        logger.debug('%s directory:%d code files.', indir, len(files))
        return files

class CodeWriter(object):
    def __init__(self, font_name='宋体', font_size=10.5, space_before=0.0, space_after=2.3, line_spacing=10.5, command_chars=None, document=None):
        self.font_name = font_name
        self.font_size = font_size
        self.space_before = space_before
        self.space_after = space_after
        self.line_spacing = line_spacing
        self.command_chars = command_chars if not command_chars else DEFAULT_COMMENT_CHARS
        self.document = Document(importlib_resources.files('swcr').joinpath('template.docx')) if not document else document

    @staticmethod
    def is_blank_line(line):
        return not bool(line.strip())

    def is_comment_line(self, line):
        return any(line.lstrip().startswith(comment_char) for comment_char in self.command_chars)

    def write_header(self, title):
        paragraph = self.document.sections[0].header.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(title)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        return self

    def check_file_encoding(self, file_path):
        """ check file encoding """
        import chardet
        with open(file_path, 'rb') as fd:
            encode_str = chardet.detect(fd.read())['encoding']
            logging.info("input_file: %s, encoding: %s", file_path, encode_str)
            return encode_str

    def write_file(self, file):
        encoding = self.check_file_encoding(file)
        print(file, encoding)
        with codecs.open(file, 'r', encoding, errors='replace') as fp:
            for line in fp:
                line = line.rstrip()
                blank_line = self.is_blank_line(line)
                comment_line = self.is_comment_line(line)
                if blank_line or comment_line:
                    continue
                # Convert the line to XML compatible text
                # line = self.convert_to_xml_text(line)
                paragraph = self.document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(self.space_before)
                paragraph.paragraph_format.space_after = Pt(self.space_after)
                paragraph.paragraph_format.line_spacing = Pt(self.line_spacing)
                run = paragraph.add_run(line)
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)
        return self

    def convert_to_xml_text(self, text):
        # Remove any NULL bytes or control characters from the text
        text = ''.join(ch for ch in text if ch.isprintable())
        return text

    def save(self, file):
        self.document.save(file)

@dataclass
class MainParams:
    title: str
    indirs: list
    exts: list
    comment_chars: list
    font_name: str
    font_size: float
    space_before: float
    space_after: float
    line_spacing: float
    excludes: list
    outfile: str
    verbose: bool

def main(main_params: MainParams):
    title = main_params.title
    indirs = main_params.indirs
    exts = main_params.exts
    comment_chars = main_params.comment_chars
    font_name = main_params.font_name
    font_size = main_params.font_size
    space_before = main_params.space_before
    space_after = main_params.space_after
    line_spacing = main_params.line_spacing
    excludes = main_params.excludes
    outfile = main_params.outfile
    verbose = main_params.verbose

    if not indirs:
        indirs = DEFAULT_INDIRS
    if not exts:
        exts = DEFAULT_EXTS
    if not comment_chars:
        comment_chars = DEFAULT_COMMENT_CHARS
    if verbose:
        logging.basicConfig(level=logging.DEBUG)

    # 第零步，把所有的路径都转换为绝对路径
    new_indirs = []
    for indir in indirs:
        new_indirs.append(abspath(indir))
    indirs = new_indirs

    excludes = del_slash(
        [abspath(exclude) for exclude in excludes] if excludes else []
    )

    # 第一步，查找代码文件
    finder = CodeFinder(exts)
    files = []
    for indir in indirs:
        for file in finder.find(indir, excludes = excludes):
            # print(file)
            files.append(file)
    print()

    # 第二步，逐个把代码文件写入到docx中
    writer = CodeWriter(
        command_chars = comment_chars,
        font_name = font_name,
        font_size = font_size,
        space_before = space_before,
        space_after = space_after,
        line_spacing = line_spacing
    )
    writer.write_header(title)
    for file in files:
        writer.write_file(file)
    writer.save(outfile)
    return 0

def parse_args():
    parser = argparse.ArgumentParser(description='Process some integers.')
    parser.add_argument('--title', type=str, help='Title for the document')
    parser.add_argument('--indirs', type=str, nargs='+', help='Input directories')
    parser.add_argument('--exts', type=str, nargs='+', help='File extensions')
    parser.add_argument('--comment_chars', type=str, nargs='+', help='Comment characters')
    parser.add_argument('--font_name', type=str, help='Font')
    parser.add_argument('--font_size', type=float, help='Font size')
    parser.add_argument('--space_before', type=float, help='space before')
    parser.add_argument('--space_after', type=float, help='space after')
    parser.add_argument('--line_spacing', type=float, help='line spacing')
    parser.add_argument('--outfile', type=str, help='Output file')

    args = parser.parse_args()
    return args

if __name__ == '__main__':
    args = parse_args()
    # print(args)
    params = MainParams(
        title = args.title,
        indirs = args.indirs,
        exts = args.exts,
        comment_chars = args.comment_chars,
        font_name = args.font_name,
        font_size = args.font_size,
        space_before = args.space_before,
        space_after = args.space_after,
        line_spacing = args.line_spacing,
        excludes = [],
        outfile = args.outfile,
        verbose = False
    )
    main(params)